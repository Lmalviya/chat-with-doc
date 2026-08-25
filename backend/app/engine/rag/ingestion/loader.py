import io
import logging
from typing import List, Dict, Any

from pypdf import PdfReader
import docx

from langchain_core.documents import Document


class DocumentLoader:
    
    @staticmethod
    def load_pdf(
        file_bytes: bytes,
        base_metadata: Dict[str, Any]
    ) -> List[Document]:
        docs = []
        reader = PdfReader(io.BytesIO(file_bytes))
        for page_idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                metadata = {
                    **base_metadata,
                    "page_number": page_idx + 1,
                    "total_pages": len(reader.pages)
                }
                docs.append(Document(
                    page_content=text.strip(),
                    metadata=metadata
                ))
        return docs

    @staticmethod
    def load_docx(
        file_bytes: bytes,
        base_metadata: Dict[str, Any]
    ) -> List[Document]:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    full_text.append(row_text)

        content = "\n\n".join(full_text)
        if not content.strip():
            return []

        return [
            Document(page_content=content, metadata=base_metadata)
        ]
    
    @staticmethod
    def load_text(file_bytes: bytes, base_metadata: Dict[str, Any]) -> List[Document]:
        try:
            content = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            logging.info("Failed to decode text file as UTF-8, trying Latin-1")
            content = file_bytes.decode("latin-1", errors="ignore")


        if not content.strip():
            return []

        return [
            Document(page_content=content.strip(), metadata=base_metadata)
        ]


    @classmethod
    def load(
        cls,
        file_bytes: bytes,
        file_name: str,
        file_type: str,
        base_metadata: Dict[str, Any] | None = None
    ) -> List[Document]:

        base_metadata = base_metadata.copy() if base_metadata else {}
        base_metadata["file_name"] = file_name
        base_metadata["file_type"] = file_type

        ext = file_name.split(".")[-1].lower() if "." in file_name else ""

        if ext == "pdf" or "pdf" in file_type:
            return cls.load_pdf(file_bytes, base_metadata)
        elif ext in ["docx", "doc"] or "word" in file_type:
            return cls.load_docx(file_bytes, base_metadata)
        else:
            return cls.load_text(file_bytes, base_metadata)

        